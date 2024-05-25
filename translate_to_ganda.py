from docx import Document
from deep_translator import GoogleTranslator

def extract_text_from_paragraph(paragraph):
    return paragraph.text

def extract_text_from_docx(doc_path):
    # Load the DOCX file
    doc = Document(doc_path)
    
    # Initialize a list to hold all text
    all_text = set()
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        all_text.add(extract_text_from_paragraph(paragraph))
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text.add(extract_text_from_paragraph(paragraph))
                    
    # remove any empty strings, numbers eg 1, 2, 3, 4, 5, 6, 7, 8, 9, 0 and ' '
    all_text = [text for text in all_text if text.strip() and not text.isdigit()]
    
    return all_text

def translate_text(text_list, target_language):
    translator = GoogleTranslator(source='en', target=target_language)
    translated_texts = [translator.translate(text) for text in text_list]
    return translated_texts

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if old_text in item.text:
                item.text = item.text.replace(old_text, new_text)

def replace_text_in_docx(doc_path, original_texts, translated_texts):
    # Load the DOCX file 
    doc = Document(doc_path)
    
    # Create a dictionary for replacements
    replacements = dict(zip(original_texts, translated_texts))
    
    # Iterate over paragraphs in the document
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            replace_text_in_paragraph(paragraph, old_text, new_text)
    
    # Iterate over tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text_in_paragraph(paragraph, old_text, new_text)
    
    # Save the modified document
    new_doc_path = doc_path.replace('.docx', '_translated.docx')
    doc.save(new_doc_path)
    return new_doc_path

# Example usage
doc_path = '3RTreview.docx'

# Extract all text groups
print('Extracting text from DOCX file...')
all_text = extract_text_from_docx(doc_path)
print(f'Extracted {len(all_text)} text groups')

# Translate all text to Luganda
print('Translating text to Luganda...')
translated_text = translate_text(all_text, 'lg')
print('Translation complete')

# Replace original text with translated text in the DOCX file
print('Replacing original text with translated text...')
new_doc_path = replace_text_in_docx(doc_path, all_text, translated_text)
print(f'Translated document saved as: {new_doc_path}')
